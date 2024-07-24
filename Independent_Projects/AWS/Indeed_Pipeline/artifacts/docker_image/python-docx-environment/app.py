import json
import pandas as pd 
import boto3
from docx import Document
from io import BytesIO
import re 

# Initialize the S3 and Glue clients
s3 = boto3.client('s3')
glue = boto3.client('glue')

def lambda_handler(event, context):
    try:
        # Print the event to see the structure when testing
        print("Main Function Initiated")
        print("Event: ", json.dumps(event, indent=2))

        # Define the bucket and the raw folder
        bucket = event['Records'][0]['s3']['bucket']['name']
        print("Bucket Name: ", bucket)
        raw_folder = 'raw/'
        processed_word_docx_folder = 'processed_word_docx/'
        raw_processed = 'raw_processed/'
        
        # List all objects in the raw folder
        response = s3.list_objects_v2(Bucket=bucket, Prefix=raw_folder)

        if 'Contents' not in response:
            raise KeyError("'Contents' key is missing in the list_objects_v2 response")

        # Initialize the continuation token for pagination
        continuation_token = None

        # Loop through all objects in the raw folder
        while True:
            if continuation_token:
                response = s3.list_objects_v2(Bucket=bucket, Prefix=raw_folder, ContinuationToken=continuation_token)
            else:
                response = s3.list_objects_v2(Bucket=bucket, Prefix=raw_folder)

            if 'Contents' not in response:
                raise KeyError("'Contents' key is missing in the list_objects_v2 response")

            # Extract keys from JSON response variable
            for obj in response['Contents']:
                key = obj['Key']
                if key.endswith('/'):  # Skip any folder-like objects
                    continue

                # Get the DOCX file from the S3 bucket
                file_response = s3.get_object(Bucket=bucket, Key=key)
                file_content = file_response['Body'].read()

                # Load the DOCX file with python-docx
                document_to_string = read_word(BytesIO(file_content))
                print("document successfully converted to string object")

                # Modify the document, split each description separately then clean using function
                separate_job_objects = document_to_string.split('\n\n--------------------') # Separate the job from a divider marker my bot set
                repaired_doc = clean_job_description_paragraph(separate_job_objects)

                # Split cleaned text object (repaired_doc), grab filename/year, and then convert to dataframe object to save as csv
                csv_split_file = repaired_doc.split('--------------------------------------------------------------------------------------')
                word_docx_title = key.split('/')[-1][:-5]
                report_year = word_docx_title[:-5].split('_')[2].split('-')[0] # extracts the year from filename
                state = word_docx_title[:-5].split('_')[1]
                docx_to_csv('s3://' + bucket + r"/" + raw_processed + word_docx_title + '.csv', csv_split_file, int(report_year), state)
                print('successfuly saved as csv')

                # Create a new document and add repaired text object to it
                cleansed_page = Document()
                cleansed_page.add_paragraph(repaired_doc)

                # Save the modified document to a BytesIO object
                output_stream = BytesIO()
                cleansed_page.save(output_stream)
                output_stream.seek(0)

                # Define new key/filepath for the modified file
                new_key = processed_word_docx_folder + key.split('/')[-1]

                # Upload the modified document to the same S3 bucket but in a different path
                s3.put_object(Bucket=bucket, Key=new_key, Body=output_stream.getvalue())

    except Exception as e:
        print(e)
        raise e
    


def clean_job_description_paragraph(job_desc_list: list) -> str:
    ''' Take dirty job description html code already separated using ('-------'), and cleans html formatting using regex expressions
    
        Args: 
        job_desc_list (list): dirty job description html code 
        
        Returns: 
        job_description_string (str): string object of cleaned job description text
        
    '''
    
    job_description_string = ''
    
    phone_number_pattern = r'\(?\b\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
    date_pattern = r'\(?\b\d{4}\)?[-.\s]?\d{2}[-.\s]?\d{2}\b'
    remove_top_html_pattern = r'</div><div id="jobDescriptionText"'
    remove_bottom_html_pattern = r'</div></div>'
    merged_tag_pattern = r'<(\w{1,2})(\w+)'
    replacement = r'\2'
    first_line_pattern = re.compile(r' class="[^"]*">\s')
    additional_line_pattern = re.compile('(.*?)class="[^>]*">')
    
    tags_to_remove = ['<div>', '</div>', '<p>', '</p>', '<br>', '</br>', '<ul>', '</ul>', '<i>', '</i>', '<b>', '</b>', '<li>', '</li>', '\n', '\n+', '<i>', "'", '<h4>', 
                        '</h4>', '</h3>', '<h3>', '<h2>', '</h2>', "’", r'/', r'\.00\b','  +']
    
    
    for job_desc_html in job_desc_list:
        
        try:
            job_desc_html_v2 = job_desc_html.split(remove_top_html_pattern)[1]
            job_desc_html_v3 = job_desc_html_v2.split(remove_bottom_html_pattern)[0]
        except:
            continue
        


        for items in tags_to_remove: 
            job_desc_html_v3 = re.sub(items, ' ', job_desc_html_v3)

        
        for regex_fliters in [first_line_pattern, phone_number_pattern, date_pattern, additional_line_pattern]:
            job_desc_html_v3 = re.sub(regex_fliters, '', job_desc_html_v3)

        
        refined_job_desc = re.sub(merged_tag_pattern, replacement, job_desc_html_v3)

        job_description_string += "".join("Job Description:  ") + refined_job_desc +  "\n\n -------------------------------------------------------------------------------------- \n\n "
        
    return job_description_string.lower()

def docx_to_csv(csv_file_path: str, input_string_list: list, report_year: int, state: str):
    ''' Takes the separated job description list and creates a dataframe to readablility in AWS Glue
    
        Args: 
        csv_file_path (str): file path to save dataframe using word_docx title rename 
        input_string_list (list): list of job descriptions split using ('------') separator
        year (int): year as integer
        state (str): state code 
    '''
    job_dataframe = pd.DataFrame()
    job_dataframe['numb_description'] = range(1, len(input_string_list))
    job_dataframe['job_description'] = input_string_list[:-1]
    job_dataframe['state'] = state
    job_dataframe['report_year'] = report_year
    job_dataframe.to_csv(path_or_buf= csv_file_path, index=False)
    
def read_word(path: str) -> str:
    ''' Creates a string object from a word document of job descriptions div containers
    
        Args: 
        path (str): file path to word document locaiton 
        
        Returns: 
        doc (str): a string made of the text of the word document 
    '''
    document = Document(path)
    
    # Initialize an empty string to hold the text block
    doc_to_text = ""

    # Iterate through each paragraph in the document
    for para in document.paragraphs:
    # Append each paragraph's text to the text block, followed by a newline character
        doc_to_text += para.text + "\n"
    
    return doc_to_text